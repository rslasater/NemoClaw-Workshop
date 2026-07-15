from __future__ import annotations

import zipfile
from pathlib import Path

from docx import Document
from docx.shared import Inches, Pt
from reportlab.lib import colors
from reportlab.lib.pagesizes import letter
from reportlab.lib.styles import getSampleStyleSheet
from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer, Table, TableStyle


REPO_ROOT = Path(__file__).resolve().parents[2]
ATTACHMENTS = REPO_ROOT / "fake-email-api" / "data" / "attachments"


def ensure_dir() -> None:
    ATTACHMENTS.mkdir(parents=True, exist_ok=True)


def add_heading(document: Document, title: str, subtitle: str) -> None:
    section = document.sections[0]
    section.top_margin = Inches(0.75)
    section.bottom_margin = Inches(0.75)
    section.left_margin = Inches(0.8)
    section.right_margin = Inches(0.8)
    heading = document.add_paragraph()
    run = heading.add_run(title)
    run.bold = True
    run.font.size = Pt(18)
    sub = document.add_paragraph()
    sub.add_run(subtitle).italic = True


def save_talking_points() -> None:
    document = Document()
    add_heading(document, "Project ATaS Talking Points", "UNCLASSIFIED - working draft")
    document.add_paragraph(
        "Purpose: provide concise, accurate language for leadership discussions and congressional staff preparation."
    )
    sections = {
        "Core Message": [
            "Project ATaS is a decision-support pilot focused on analyst workflow acceleration.",
            "Current validation supports continued evaluation, not broad operational deployment.",
            "Urban test results require careful caveating around false positives and edge cases.",
        ],
        "Use With Caution": [
            "Avoid describing the system as autonomous or fully mission-ready.",
            "Do not cite unreviewed appendix material or draft slides outside the approved package.",
            "Route security or distribution questions to the JIAITF security office.",
        ],
        "Likely Follow-Up": [
            "What changed between validation report versions?",
            "How are edge-case failures being tracked and remediated?",
            "Who has authority to approve broader use?",
        ],
    }
    for title, bullets in sections.items():
        document.add_heading(title, level=2)
        for bullet in bullets:
            document.add_paragraph(bullet, style="List Bullet")
    document.save(ATTACHMENTS / "ATaS_Talking_Points.docx")


def save_congressional_questions() -> None:
    document = Document()
    add_heading(document, "Expected Congressional Questions", "UNCLASSIFIED - staff preparation")
    table = document.add_table(rows=1, cols=3)
    table.style = "Table Grid"
    headers = table.rows[0].cells
    headers[0].text = "Topic"
    headers[1].text = "Likely Question"
    headers[2].text = "Prep Note"
    rows = [
        ("Validation", "What evidence supports the current readiness claim?", "Anchor on v18 results and identify caveats."),
        ("Urban Testing", "What failed in dense urban scenarios?", "Discuss false positives without overstating impact."),
        ("Oversight", "Who approved the briefing package?", "Confirm distribution path before answering."),
        ("Security", "Was any sensitive material mishandled?", "Refer suspected spillage to Security."),
        ("Funding", "What does the team need next?", "Tie requests to retest and governance needs."),
    ]
    for row in rows:
        cells = table.add_row().cells
        for cell, text in zip(cells, row):
            cell.text = text
    document.add_paragraph()
    document.add_paragraph(
        "Note: This document is synthetic workshop material. It intentionally avoids real operational details."
    )
    document.save(ATTACHMENTS / "Congressional_Questions.docx")


def pdf_story(filename: str, title: str, subtitle: str, rows: list[tuple[str, str]], note: str) -> None:
    styles = getSampleStyleSheet()
    doc = SimpleDocTemplate(
        str(ATTACHMENTS / filename),
        pagesize=letter,
        rightMargin=54,
        leftMargin=54,
        topMargin=54,
        bottomMargin=54,
    )
    story = [
        Paragraph(title, styles["Title"]),
        Paragraph(subtitle, styles["Italic"]),
        Spacer(1, 18),
    ]
    table_data = [["Section", "Summary"], *rows]
    table = Table(table_data, colWidths=[150, 330])
    table.setStyle(TableStyle([
        ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#1F4E78")),
        ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
        ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
        ("GRID", (0, 0), (-1, -1), 0.25, colors.HexColor("#B7C9D9")),
        ("VALIGN", (0, 0), (-1, -1), "TOP"),
        ("ROWBACKGROUNDS", (0, 1), (-1, -1), [colors.white, colors.HexColor("#F5F8FB")]),
        ("LEFTPADDING", (0, 0), (-1, -1), 8),
        ("RIGHTPADDING", (0, 0), (-1, -1), 8),
        ("TOPPADDING", (0, 0), (-1, -1), 6),
        ("BOTTOMPADDING", (0, 0), (-1, -1), 6),
    ]))
    story.extend([table, Spacer(1, 16), Paragraph(note, styles["BodyText"])])
    doc.build(story)


def save_pdfs() -> None:
    pdf_story(
        "ATaS_Validation_Report_v17.pdf",
        "ATaS Validation Report v17",
        "UNCLASSIFIED - preliminary validation summary",
        [
            ("Scope", "Covers September desktop and limited field validation runs."),
            ("Result", "Shows promise but leaves urban edge cases open for retest."),
            ("Primary Gap", "False positives remain above desired threshold in dense urban scenarios."),
            ("Recommendation", "Use guarded language until v18 retest results are confirmed."),
        ],
        "Prepared for workshop use. This synthetic report does not contain real operational information.",
    )
    pdf_story(
        "ATaS_Validation_Report_v18.pdf",
        "ATaS Validation Report v18",
        "UNCLASSIFIED - updated validation summary",
        [
            ("Scope", "Adds late September urban retest results and revised analyst workflow measures."),
            ("Result", "Supports a conditional pass with explicit caveats."),
            ("Improvement", "False negatives improved, but false positives remain a leadership concern."),
            ("Recommendation", "Brief as a controlled pilot with further remediation required."),
        ],
        "Prepared for workshop use. This synthetic report does not contain real operational information.",
    )
    pdf_story(
        "VA_Loan_Preapproval.pdf",
        "VA Loan Preapproval Letter",
        "Personal finance attachment - synthetic workshop artifact",
        [
            ("Applicant", "James Maddox"),
            ("Status", "Preapproved subject to final income, credit, and property review."),
            ("Amount", "Up to USD 885,000"),
            ("Expiration", "2026-11-15"),
        ],
        "This training artifact is fictional and should not be used for lending decisions.",
    )
    pdf_story(
        "Inspection_Report.pdf",
        "Residential Inspection Summary",
        "Synthetic home inspection attachment",
        [
            ("Property", "San Diego area single-family residence"),
            ("Roof", "Serviceable, monitor minor flashing wear."),
            ("Electrical", "Panel labeled; recommend GFCI check in garage."),
            ("Plumbing", "No active leaks observed during walkthrough."),
            ("Follow-Up", "Request seller clarification on HVAC service history."),
        ],
        "This synthetic report supports personal-noise email threads in the workshop mailbox.",
    )


def save_calendar() -> None:
    text = "\r\n".join([
        "BEGIN:VCALENDAR",
        "VERSION:2.0",
        "PRODID:-//JIAITF//Project ATaS Workshop//EN",
        "BEGIN:VEVENT",
        "UID:atas-briefing-agenda-20261006@jiaitf.example",
        "DTSTAMP:20260929T120000Z",
        "DTSTART:20261006T130000Z",
        "DTEND:20261006T143000Z",
        "SUMMARY:Project ATaS Congressional Briefing Prep",
        "LOCATION:JIAITF Main Conference Room",
        "DESCRIPTION:Review agenda, expected questions, validation caveats, and approved speaking roles.",
        "END:VEVENT",
        "END:VCALENDAR",
        "",
    ])
    (ATTACHMENTS / "ATaS_Briefing_Agenda.ics").write_text(text, encoding="utf-8", newline="")


def save_v12_deck() -> None:
    source = ATTACHMENTS / "ATaS_Congressional_Brief_v13.pptx"
    target = ATTACHMENTS / "ATaS_Congressional_Brief_v12.pptx"
    with zipfile.ZipFile(source, "r") as zin, zipfile.ZipFile(target, "w", compression=zipfile.ZIP_DEFLATED) as zout:
        for item in zin.infolist():
            data = zin.read(item.filename)
            if item.filename == "ppt/slides/slide14.xml":
                text = data.decode("utf-8")
                text = text.replace("SECRET//NOFORN", "UNCLASSIFIED")
                text = text.replace("erroneously embedded", "draft appendix")
                data = text.encode("utf-8")
            zout.writestr(item, data)


def main() -> None:
    ensure_dir()
    save_talking_points()
    save_congressional_questions()
    save_pdfs()
    save_calendar()
    save_v12_deck()


if __name__ == "__main__":
    main()
